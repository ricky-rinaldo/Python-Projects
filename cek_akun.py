# template_pertanyaan.py

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor


def set_font_style(run, size, bold=False, color=None):
    """Fungsi pembantu untuk mengatur gaya font."""
    font = run.font
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = color


def generate_template_document():
    """
    Membuat dokumen Word yang berisi template pertanyaan analisis hook.
    """
    document = Document()

    # 1. JUDUL DOKUMEN
    document.add_heading('TEMPLATE PERTANYAAN ANALISIS HOOK VIDEO', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    set_font_style(p.add_run('Gunakan template ini untuk meminta kode analisis klip hook video ke Gemini AI.'), 11)

    document.add_paragraph().add_run().add_break()

    # 2. INSTRUKSI PENGGUNAAN
    document.add_heading('Petunjuk Penggunaan', level=1)

    instructions = [
        "1. Isi semua bagian yang ditandai dengan [INPUT ANDA: ...].",
        "2. Pastikan START_OFFSET_SECONDS diisi dengan benar (dalam detik) untuk menghindari deteksi pada bumper/cuplikan awal.",
        "3. Kriteria Skor Hook (4-5) dan Durasi Klip Minimal (misalnya: 15 detik) bersifat wajib.",
        "4. Salin (copy) seluruh teks dari [PERMINTAAN MULAI DI SINI] hingga [PERMINTAAN BERAKHIR DI SINI] dan tempel (paste) ke jendela chat AI."
    ]
    for inst in instructions:
        document.add_paragraph(inst, style='List Bullet')

    document.add_paragraph().add_run().add_break()

    # 3. KOTAK TEMPLATE
    document.add_heading('TEMPLATE PERMINTAAN', level=1)

    # Garis Pembatas Awal
    p_start = document.add_paragraph()
    set_font_style(p_start.add_run('[PERMINTAAN MULAI DI SINI]'), 11, bold=True, color=RGBColor(0x00, 0x00, 0xFF))

    # Tambahkan seluruh template pertanyaan sebagai satu blok teks
    template_text = """
"Tolong buatkan kode Python untuk dokumentasi dan pemformatan analisis klip hook ke dalam file Word (.docx), menggunakan kriteria sensitivitas (Skor 4-5) dan format yang sama seperti sebelumnya.

Output file harus disimpan di folder F:\\Ricks\\Clipper Video\\Video\\Hook.

Berikut data video dan analisis hook yang sudah saya kurasi:

- LINK_YOUTUBE: [INPUT ANDA: URL YouTube]
- JUDUL_VIDEO: [INPUT ANDA: Judul Lengkap Video]
- DURASI_MINIMAL_KLIP: [INPUT ANDA: Durasi Minimal Klip dalam Detik, misalnya: 15]
- START_OFFSET_SECONDS: [INPUT ANDA: Waktu Awal Deteksi Hook (dalam Detik) setelah Bumper/Cuplikan, misalnya: 30]

--- DATA HOOK KLIP TERBAIK (Target: 12 sampai 15 Hook, Skor 4-5) ---
# Format: (Skor Hook, Start Time [HH:MM:SS], End Time [HH:MM:SS], Durasi Klip [XX detik], Kalimat Hook)

data_hook = [
    # Minimal 12, Maksimal 15 Hook Kurasi Anda:
    [INPUT ANDA: Data Hook Kurasi Anda (Wajib Skor 4 atau 5)],
    (5, "00:01:30", "00:01:45", "15 detik", "Contoh kalimat hook pertama."),
    (4, "00:03:00", "00:03:25", "25 detik", "Contoh kalimat hook kedua."),
    # ... Lanjutkan hingga 12-15 klip kurasi Anda di sini ...
]

Buatkan seluruh kode Python-nya."
"""

    # Menambahkan teks template
    p_template = document.add_paragraph()
    template_run = p_template.add_run(template_text.strip())
    set_font_style(template_run, 10)

    # Garis Pembatas Akhir
    p_end = document.add_paragraph()
    set_font_style(p_end.add_run('[PERMINTAAN BERAKHIR DI SINI]'), 11, bold=True, color=RGBColor(0x00, 0x00, 0xFF))

    # 4. SIMPAN DOKUMEN
    output_directory = r"F:\Ricks\Clipper Video\Video\Hook"
    output_filename = "template pertanyaan.docx"

    # Pastikan direktori ada
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
        print(f"⚠️ Direktori dibuat: {output_directory}")

    # Gabungkan path
    output_filepath = os.path.join(output_directory, output_filename)

    document.save(output_filepath)

    print("-" * 50)
    print(f"✅ Dokumen template berhasil dibuat!")
    print(f"📂 File disimpan sebagai: {os.path.abspath(output_filepath)}")
    print("-" * 50)


# ==============================================================================
# EKSEKUSI PROGRAM
# ==============================================================================
if __name__ == "__main__":
    generate_template_document()