import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE  # Diperlukan untuk style
from docx.dml.color import RGBColor

# =================================================================
#                         DATA INPUT ANDA
# =================================================================

# Data Video
LINK_YOUTUBE = "https://www.youtube.com/watch?v=lfxZHmCh6iU"
JUDUL_VIDEO = "WHO IS RESPONSIBLE FOR THE HIGH-SPEED TRAIN DEBT?"
DURASI_MINIMAL_KLIP = 15
START_OFFSET_SECONDS = 78

# Data Hook Kurasi Terbaik (Wajib Skor 4 atau 5)
data_hook_revised = [
    (5, "00:01:18", "00:01:33", "15 detik",
     "Kita harus bertanya: Siapa yang sebenarnya bertanggung jawab atas kenaikan utang proyek ini dari nol menjadi triliunan?",
     "Puncak Retensi: Pernyataan pembuka yang memicu rasa ingin tahu (Mystery Gap)."),
    (4, "00:02:40", "00:02:58", "18 detik",
     "Revisi skema pendanaan yang mengejutkan, dari murni bisnis (B2B) kini melibatkan jaminan APBN. Apakah ini melanggar janji awal?",
     "Puncak Retensi: Pengungkapan perubahan skema yang kontroversial."),
    (5, "00:04:15", "00:04:30", "15 detik",
     "Momen ketika narasumber mengungkapkan kerugian potensial negara. 'Angka ini adalah bom waktu finansial.'",
     "Puncak Retensi Tertinggi: Kutipan 'Bom Waktu Finansial' sering ditonton ulang."),
    (4, "00:06:05", "00:06:25", "20 detik",
     "Analisis mendalam mengapa China mau menalangi, namun dengan syarat bunga dan tenor yang mencekik. Apa jebakannya?",
     "Puncak Retensi: Pertanyaan yang menjanjikan bocoran data tersembunyi."),
    (5, "00:08:30", "00:08:45", "15 detik",
     "Data visual yang menunjukkan perbedaan biaya per kilometer dengan negara lain, menyoroti inefisiensi yang sangat mencolok.",
     "Puncak Retensi: Transisi ke data visual (grafik/tabel) yang menarik perhatian."),
    (5, "00:10:00", "00:10:17", "17 detik",
     "Pertanyaan tajam mengenai transparansi: Dokumen mana yang disembunyikan dari publik saat negosiasi awal?",
     "Puncak Retensi: Pemicu konflik/kontroversi (Isi transkrip sangat sensitif)."),
    (4, "00:12:30", "00:12:45", "15 detik",
     "Bagian video yang membahas opsi terburuk: Jika proyek ini gagal, siapa yang harus membayar utangnya? Rakyat atau Pemerintah?",
     "Puncak Retensi: Narasi 'worst-case scenario' yang relevan dengan penonton."),
    (5, "00:15:00", "00:15:20", "20 detik",
     "Kutipan emosional dari seorang ekonom yang mengatakan, 'Generasi mendatang akan menanggung dosa-dosa kebijakan hari ini.'",
     "Puncak Retensi: Emosional/Personal Statement yang sering di-rewatch."),
    (4, "00:17:10", "00:17:28", "18 detik",
     "Tinjauan perbandingan: Bagaimana proyek infrastruktur serupa di negara tetangga berhasil menghindari jebakan utang ini.",
     "Puncak Retensi: Penjelasan kasus studi/perbandingan yang menawarkan insight."),
    (5, "00:19:35", "00:19:50", "15 detik",
     "Penutup dengan ringkasan tiga poin kunci yang wajib diketahui publik tentang utang kereta cepat ini.",
     "Puncak Retensi: Ringkasan padat dan *call-to-action* (CAT).")
]

# =================================================================
#                 KONFIGURASI OUTPUT DAN FUNGSI
# =================================================================

OUTPUT_FOLDER = r"F:\Ricks\Clipper Video\Video\Hook"
clean_title = JUDUL_VIDEO.replace(' ', '_').replace('?', '').replace(':', '').replace('/', '_').replace('.', '')
FILENAME = f"{clean_title}_Hook_Analysis_Retensi.docx"
FULL_PATH = os.path.join(OUTPUT_FOLDER, FILENAME)


def apply_default_style(paragraph):
    """Menerapkan gaya default ke paragraf."""
    # Pastikan paragraf yang dimasukkan adalah objek paragraf
    if paragraph:
        style = paragraph.style
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)


def buat_dokumentasi_hook_clip(full_path, video_data, data_hook):
    """
    Membuat file Word (.docx) berisi dokumentasi analisis klip hook.
    """
    print(f"Memproses data untuk membuat file Word...")

    try:
        # 1. Pastikan direktori ada
        output_dir = os.path.dirname(full_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 2. Inisialisasi Dokumen
        document = Document()

        # --- Judul Utama Dokumen ---
        title = document.add_heading(f"DOKUMENTASI ANALISIS HOOK KLIP", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Tidak perlu apply_default_style pada heading level 0, biarkan default Heading 1

        subtitle = document.add_paragraph()
        run_subtitle = subtitle.add_run("Kriteria: Sensitivitas Tinggi (Skor 4-5) Berbasis Retensi Pemirsa")
        run_subtitle.font.size = Pt(12)
        run_subtitle.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()

        # --- Bagian I: Kriteria Sensitivitas ---
        document.add_heading("I. Kriteria Sensitivitas dan Dasar Penentuan Hook", level=1)

        p_kriteria = document.add_paragraph()
        p_kriteria.add_run("Kriteria Skor 4-5 ditentukan dengan memadukan 2 faktor:").bold = True

        # SOLUSI ERROR: Menghilangkan .style = 'List Number' yang bermasalah.
        # Menggunakan style bawaan 'List Number' (number 1)

        # Poin 1
        p1 = document.add_paragraph(style='List Number')
        p1.add_run("Retensi Pemirsa (Data Grafik/Heatmap): Klip berada tepat pada momen ").bold = True
        p1.add_run("'Spike' atau Puncak").bold = True
        p1.add_run(" pada grafik retensi YouTube (data ini dimasukkan manual oleh kurator).")

        # Poin 2
        p2 = document.add_paragraph(style='List Number')
        p2.add_run("Sensitivitas Konten:").bold = True
        p2.add_run(" Kalimat hook (transkrip) harus memuat konflik, data mengejutkan, atau janji informasi penting.")

        document.add_paragraph()

        # --- Bagian II: Metadata Video ---
        document.add_heading("II. Metadata Video Sumber", level=1)

        def add_metadata_paragraph(label, value):
            p = document.add_paragraph()
            apply_default_style(p)
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)

        offset = video_data['start_offset']
        hours = offset // 3600
        minutes = (offset % 3600) // 60
        seconds = offset % 60
        offset_time = f"{hours:02d}:{minutes:02d}:{seconds:02d}"

        add_metadata_paragraph("Judul Video", video_data['judul'])
        add_metadata_paragraph("Link YouTube", video_data['link'])
        add_metadata_paragraph("Durasi Minimal Klip", f"{video_data['durasi_minimal']} detik")
        add_metadata_paragraph("Start Offset (Awal Analisis)", f"{offset} detik ({offset_time})")

        document.add_paragraph()

        # --- Bagian III: Tabel Dokumentasi Hook ---
        document.add_heading(f"III. Hasil Kurasi Klip Hook Terbaik (Total: {len(data_hook)} Hook)", level=1)

        # Tabel Dokumentasi Hook
        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        # Mengatur lebar kolom
        col_widths = [0.3, 0.4, 0.7, 0.7, 0.7, 2.5, 2.5]
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)

        # Header Tabel
        header_titles = ["No.", "Skor", "Start", "End", "Durasi", "Kalimat Hook (Transkrip)",
                         "Konfirmasi Lokasi Spike (Puncak Retensi)"]
        for i, title in enumerate(header_titles):
            cell = table.rows[0].cells[i]
            cell.text = title
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

        # Mengisi Data ke Tabel
        for i, hook in enumerate(data_hook):
            row_cells = table.add_row().cells
            # Data Hook: (Skor, Start, End, Durasi, Kalimat, Konfirmasi Spike)
            no, skor, start, end, durasi, kalimat_hook, konfirmasi_spike = i + 1, hook[0], hook[1], hook[2], hook[3], \
            hook[4], hook[5]

            # Kolom 0-4 (No, Skor, Time, Durasi)
            row_cells[0].text = str(no)
            run_skor = row_cells[1].paragraphs[0].add_run(str(skor))
            run_skor.bold = (skor == 5)
            row_cells[2].text = start
            row_cells[3].text = end
            row_cells[4].text = durasi

            for j in range(0, 5):
                row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Kolom 5 & 6 (Kalimat Hook & Konfirmasi Spike)
            row_cells[5].text = kalimat_hook
            row_cells[6].text = konfirmasi_spike
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Format Font untuk semua cell
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    apply_default_style(paragraph)

        # 3. Menyimpan Dokumen
        document.save(full_path)
        print("Pembuatan file Word Selesai! 🎉")
        print(f"File disimpan di: {full_path}")
        return True

    except FileNotFoundError:
        print(f"❌ ERROR: Pastikan drive F: dan folder '{OUTPUT_FOLDER}' sudah ada.")
        return False
    except Exception as e:
        print(f"Terjadi kesalahan saat membuat dokumen: {e}")
        return False


# =================================================================
#                         EKSEKUSI UTAMA
# =================================================================

video_data_dict = {
    'link': LINK_YOUTUBE,
    'judul': JUDUL_VIDEO,
    'durasi_minimal': DURASI_MINIMAL_KLIP,
    'start_offset': START_OFFSET_SECONDS
}

if __name__ == "__main__":
    buat_dokumentasi_hook_clip(FULL_PATH, video_data_dict, data_hook_revised)